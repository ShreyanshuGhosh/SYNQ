"""Celery tasks — the file parsing pipeline.

Implements SYNQ_STRUCT §"Tier 6 — File Pipeline":
    parse PDFs (pdfplumber) | OCR images (tesseract) | parse DOCX |
    vision description (separate cheap model) | chunk long docs

On success: ``files.parse_status='done'`` and the relevant columns are
populated. On failure: ``files.parse_status='failed'`` and the row's
``error`` JSONB carries ``{message, kind, ...}`` so the chat UI can
render a red chip.

Originals are NEVER deleted from S3. The worker only READS from S3.
"""

from __future__ import annotations

import io
import logging
import traceback
from typing import Any
from uuid import UUID

from app.storage import download_bytes, key_from_storage_url
from app.workers.celery_app import celery_app
from app.workers.db_sync import sync_session
from app.workers.description import describe_image

logger = logging.getLogger(__name__)


# Roughly 4 chars per token, matching the orchestrator's `_char_estimate`
# heuristic. Cheap, conservative, never under-chunks.
_CHARS_PER_TOKEN = 4


@celery_app.task(name="app.workers.tasks.heartbeat")
def heartbeat() -> str:
    """Beat-driven no-op. Confirms the schedule + broker are live."""
    return "ok"


@celery_app.task(name="app.workers.tasks.parse_file", bind=True, max_retries=2)
def parse_file(self: Any, file_id: str, ext: str) -> dict[str, Any]:
    """Parse the file identified by `file_id`. Updates parse_status.

    Dispatch by extension is intentional and lives ONLY in this function.
    Adapters do not care how a description was produced — they only
    consume the resulting columns.
    """
    from app.orm import File  # local import: celery loads tasks at boot

    uid = UUID(file_id)
    try:
        with sync_session() as session:
            row = session.get(File, uid)
            if row is None:
                logger.warning("parse_file: file_id=%s gone before parse", uid)
                return {"status": "missing"}
            key = key_from_storage_url(row.storage_url)
            mime = row.mime_type or ""
            raw = download_bytes(key)

            update: dict[str, Any] = {}
            if ext == "pdf":
                update = _parse_pdf(raw)
            elif ext == "docx":
                update = _parse_docx(raw)
            elif ext in {"png", "jpg", "jpeg", "webp"}:
                update = _parse_image(raw, mime)
            elif ext in {"txt", "md"}:
                update = _parse_text(raw)
            else:
                raise ValueError(f"unsupported extension: {ext!r}")

            # Apply chunking ONLY for long extracted text (docs and OCR).
            text = update.get("extracted_text") or ""
            chunks = _chunk_text(text) if text else []
            if chunks:
                update["chunks"] = chunks

            for k, v in update.items():
                setattr(row, k, v)
            row.parse_status = "done"
            row.error = None
            logger.info(
                "parse_file: file_id=%s ext=%s ok (text=%dchars chunks=%d desc=%s)",
                uid,
                ext,
                len(text),
                len(chunks),
                "yes" if row.description else "no",
            )
            return {"status": "done", "chunks": len(chunks)}
    except Exception as exc:  # noqa: BLE001 — we mark failure and reraise
        logger.exception("parse_file: file_id=%s failed", uid)
        try:
            with sync_session() as session:
                row = session.get(File, uid)
                if row is not None:
                    row.parse_status = "failed"
                    row.error = {
                        "message": str(exc) or exc.__class__.__name__,
                        "kind": exc.__class__.__name__,
                        "traceback": traceback.format_exc(limit=4),
                    }
        except Exception:
            logger.exception("parse_file: also failed to record failure row=%s", uid)
        # Surface the error to Celery so the dashboard / future
        # retries can observe it. Don't auto-retry by default — most
        # parse failures (bad PDF, corrupt image) are deterministic.
        raise


# ── Per-type parsers ────────────────────────────────────────────────────


def _parse_pdf(raw: bytes) -> dict[str, Any]:
    """pdfplumber for text extraction. We do NOT call OCR on PDFs in
    Phase 3 — scanned-PDF OCR is a separate Phase 6 problem."""
    import pdfplumber

    pages_text: list[str] = []
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            pages_text.append(page.extract_text() or "")
    text = "\n\n".join(pages_text).strip()
    return {"extracted_text": text}


def _parse_docx(raw: bytes) -> dict[str, Any]:
    import docx

    document = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in document.paragraphs if p.text]
    # Tables are flattened row-by-row; cell separator = tab.
    for table in document.tables:
        for r in table.rows:
            parts.append("\t".join(cell.text for cell in r.cells))
    return {"extracted_text": "\n".join(parts).strip()}


def _parse_image(raw: bytes, mime: str) -> dict[str, Any]:
    """Tesseract OCR + vision description.

    Tesseract is best-effort: if it's not installed locally, we still
    surface a description from the vision model. The two columns are
    independent so a chat against a text-only model can fall back to
    `description` while a vision-capable target uses the raw bytes.
    """
    update: dict[str, Any] = {}

    # OCR (optional — depends on local tesseract install).
    try:
        from PIL import Image
        import pytesseract

        img = Image.open(io.BytesIO(raw))
        ocr_text = (pytesseract.image_to_string(img) or "").strip()
        if ocr_text:
            update["extracted_text"] = ocr_text
    except Exception:
        logger.info("parse_image: tesseract unavailable or failed; continuing")

    # Vision description (independent cheap model).
    description = describe_image(raw, mime)
    if description:
        update["description"] = description
    return update


def _parse_text(raw: bytes) -> dict[str, Any]:
    # UTF-8 with replacement so we never blow up on malformed bytes;
    # md and txt are both treated as plain text.
    text = raw.decode("utf-8", errors="replace").strip()
    return {"extracted_text": text}


# ── Chunking ────────────────────────────────────────────────────────────


def _chunk_text(text: str) -> list[dict[str, Any]]:
    """Split text into ~chunk_target_tokens chunks IF total is long.

    Phase 3 is naive: paragraph-aware splitting with no overlap. Phase 4
    will swap in semantic chunking; the JSONB shape stays
    `[{chunk_id, text, page}]` so the upgrade is in-place.
    """
    from app.config import settings

    if not text:
        return []
    total_tokens_estimate = len(text) // _CHARS_PER_TOKEN
    if total_tokens_estimate < settings.chunk_trigger_tokens:
        return []
    target_chars = settings.chunk_target_tokens * _CHARS_PER_TOKEN

    chunks: list[dict[str, Any]] = []
    buf: list[str] = []
    buf_len = 0
    chunk_id = 0
    for paragraph in text.split("\n\n"):
        p_len = len(paragraph)
        if buf_len + p_len > target_chars and buf:
            chunks.append(
                {
                    "chunk_id": chunk_id,
                    "text": "\n\n".join(buf),
                    "page": None,
                }
            )
            chunk_id += 1
            buf = [paragraph]
            buf_len = p_len
        else:
            buf.append(paragraph)
            buf_len += p_len + 2
    if buf:
        chunks.append(
            {"chunk_id": chunk_id, "text": "\n\n".join(buf), "page": None}
        )
    return chunks
